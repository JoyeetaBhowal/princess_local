from __future__ import annotations

import csv
import hashlib
import re
import shutil
import sqlite3
from dataclasses import dataclass
from pathlib import Path


@dataclass
class KnowledgeDocument:
    id: int
    title: str
    source_path: str
    stored_path: str
    file_type: str
    status: str
    chunk_count: int
    created_at: str


class KnowledgeLibrary:
    """Local document store with SQLite FTS search."""

    def __init__(self, base_dir: str | Path = "data"):
        self.base_dir = Path(base_dir)
        self.knowledge_dir = self.base_dir / "knowledge"
        self.index_dir = self.base_dir / "index"
        self.db_path = self.index_dir / "knowledge.db"
        self.knowledge_dir.mkdir(parents=True, exist_ok=True)
        self.index_dir.mkdir(parents=True, exist_ok=True)
        self._init_db()

    def _init_db(self):
        with sqlite3.connect(self.db_path) as conn:
            conn.execute("""
                CREATE TABLE IF NOT EXISTS documents (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    title TEXT NOT NULL,
                    source_path TEXT NOT NULL,
                    stored_path TEXT NOT NULL UNIQUE,
                    file_type TEXT NOT NULL,
                    status TEXT NOT NULL,
                    chunk_count INTEGER NOT NULL DEFAULT 0,
                    created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP
                )
            """)
            conn.execute("""
                CREATE VIRTUAL TABLE IF NOT EXISTS knowledge_chunks USING fts5(
                    document_id UNINDEXED,
                    chunk_index UNINDEXED,
                    content
                )
            """)

    def list_documents(self) -> list[KnowledgeDocument]:
        with sqlite3.connect(self.db_path) as conn:
            conn.row_factory = sqlite3.Row
            rows = conn.execute("""
                SELECT id, title, source_path, stored_path, file_type, status, chunk_count, created_at
                FROM documents
                ORDER BY created_at DESC
            """).fetchall()
        return [KnowledgeDocument(**dict(row)) for row in rows]

    def index_file(self, file_path: str | Path) -> KnowledgeDocument:
        source = Path(file_path)
        if not source.exists() or not source.is_file():
            raise FileNotFoundError(str(source))

        stored = self._copy_to_library(source)
        file_type = stored.suffix.lower().lstrip(".") or "file"
        title = source.name

        try:
            text = self._extract_text(stored)
            chunks = self._chunk_text(text)
            status = "Indexed" if chunks else "Failed"
        except Exception:
            chunks = []
            status = "Failed"

        with sqlite3.connect(self.db_path) as conn:
            existing = conn.execute(
                "SELECT id FROM documents WHERE stored_path = ?",
                (str(stored),),
            ).fetchone()
            if existing:
                document_id = int(existing[0])
                conn.execute("DELETE FROM knowledge_chunks WHERE document_id = ?", (document_id,))
                conn.execute("""
                    UPDATE documents
                    SET title = ?, source_path = ?, file_type = ?, status = ?, chunk_count = ?
                    WHERE id = ?
                """, (title, str(source), file_type, status, len(chunks), document_id))
            else:
                cursor = conn.execute("""
                    INSERT INTO documents (title, source_path, stored_path, file_type, status, chunk_count)
                    VALUES (?, ?, ?, ?, ?, ?)
                """, (title, str(source), str(stored), file_type, status, len(chunks)))
                document_id = int(cursor.lastrowid)

            conn.executemany(
                "INSERT INTO knowledge_chunks (document_id, chunk_index, content) VALUES (?, ?, ?)",
                [(document_id, index, chunk) for index, chunk in enumerate(chunks)],
            )

        return next(doc for doc in self.list_documents() if doc.id == document_id)

    def search(self, query: str, limit: int = 5) -> list[dict]:
        match = self._fts_query(query)
        if not match:
            return []
        try:
            with sqlite3.connect(self.db_path) as conn:
                conn.row_factory = sqlite3.Row
                rows = conn.execute("""
                    SELECT
                        d.title,
                        d.stored_path,
                        c.chunk_index,
                        snippet(knowledge_chunks, 2, '[', ']', '...', 24) AS snippet
                    FROM knowledge_chunks c
                    JOIN documents d ON d.id = c.document_id
                    WHERE knowledge_chunks MATCH ?
                    ORDER BY bm25(knowledge_chunks)
                    LIMIT ?
                """, (match, limit)).fetchall()
            return [dict(row) for row in rows]
        except sqlite3.OperationalError:
            return []

    def _copy_to_library(self, source: Path) -> Path:
        digest = hashlib.sha1(str(source.resolve()).encode("utf-8")).hexdigest()[:10]
        target = self.knowledge_dir / f"{source.stem}-{digest}{source.suffix.lower()}"
        if not target.exists() or source.stat().st_mtime > target.stat().st_mtime:
            shutil.copy2(source, target)
        return target

    def _extract_text(self, path: Path) -> str:
        suffix = path.suffix.lower()
        if suffix in {".txt", ".md", ".markdown", ".log"}:
            return path.read_text(encoding="utf-8", errors="ignore")
        if suffix == ".csv":
            return self._extract_csv(path)
        if suffix == ".pdf":
            return self._extract_pdf(path)
        if suffix == ".docx":
            return self._extract_docx(path)
        raise ValueError(f"Unsupported file type: {suffix or path.name}")

    def _extract_csv(self, path: Path) -> str:
        rows = []
        with path.open("r", encoding="utf-8", errors="ignore", newline="") as handle:
            reader = csv.reader(handle)
            for row in reader:
                rows.append(" | ".join(cell.strip() for cell in row if cell.strip()))
        return "\n".join(rows)

    def _extract_pdf(self, path: Path) -> str:
        try:
            from pypdf import PdfReader
        except Exception:
            try:
                from PyPDF2 import PdfReader
            except Exception as exc:
                raise RuntimeError("PDF support needs pypdf or PyPDF2 installed.") from exc
        reader = PdfReader(str(path))
        return "\n".join(page.extract_text() or "" for page in reader.pages)

    def _extract_docx(self, path: Path) -> str:
        try:
            import docx
        except Exception as exc:
            raise RuntimeError("DOCX support needs python-docx installed.") from exc
        document = docx.Document(str(path))
        return "\n".join(paragraph.text for paragraph in document.paragraphs)

    def _chunk_text(self, text: str, size: int = 1200, overlap: int = 160) -> list[str]:
        normalized = text.replace("\ufeff", " ")
        normalized = re.sub(r"\s+", " ", normalized).strip()
        if not normalized:
            return []
        chunks = []
        start = 0
        while start < len(normalized):
            end = min(start + size, len(normalized))
            chunk = normalized[start:end].strip()
            if chunk:
                chunks.append(chunk)
            if end == len(normalized):
                break
            start = max(0, end - overlap)
        return chunks

    def _fts_query(self, query: str) -> str:
        terms = re.findall(r"[A-Za-z0-9_]{2,}", query.lower())
        return " OR ".join(f'"{term}"' for term in terms[:8])


knowledge_library = KnowledgeLibrary()
